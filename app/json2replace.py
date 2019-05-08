# 使用json替换docx/doc的数据
# @chenye 2019/2/2

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import os


'''
json dumps把数据类型转换成字符串 dump把数据类型转换成字符串并存储在文件中  loads把字符串转换成数据类型  load把文件打开从字符串转换成数据类型
'''

def Replace(document, data1, data2):
    # 段落文本替换
    for a in document.paragraphs:
        for b in a.runs:
            # print(b.text) # 打印替换日志
            if b.text == data1:
                b.text = data2

    # 表格文本替换
    for table in document.tables:
        for row, obj_row in enumerate(table.rows):
            for col, cell in enumerate(obj_row.cells):
                for a in cell.paragraphs:
                    for b in a.runs:
                        # print(b.text) # 打印替换日志
                        if b.text == data1:
                            b.text = data2

def Readjson(filename):
    data = {}
    fp = open(filename, encoding='utf-8')
    js = json.load(fp)
    for i in js['result']['data'].keys():
        # 取出的键值对的值必须是str类型
        if type(js['result']['data'][i]) == type('str'):
            data[i] = js['result']['data'][i]   # 保存键值对
        else:
            print(js['result']['data'][i],' type is ',type(js['result']['data'][i]),'>>>类型不匹配（必须是str）')
    return data

def main():
    # 提取数据
    data1 = Readjson('users_show.json')
    data2 = Readjson('users_show2.json')
    # 打开docx模板文件
    document = Document('table1.docx')
    # 遍历json提取出来的字典
    for key in data1.keys():
        Replace(document, data1[key], data2[key])
    # 保存新文件
    document.save('table2.docx')

def autoproc(data2,id_):
    data1 = Readjson(os.getcwd()+'\\app\\jsons\\'+str(id_)+'.json')
    document = Document(os.getcwd()+'\\app\\docxs\\'+str(id_)+'.docx')
    for key in data1.keys():
        Replace(document, data1[key], data2[key])
    document.save(os.getcwd()+'\\app\\newdocx\\'+str(id_)+'.docx')
    #文件名暂时没有取
    return('执行成功')

#print(type(Readjson(os.getcwd()+'\\app\\jsons\\'+str(1)+'.json')))

#if __name__ == '__main__':
#    main()

'''
1. 像申请理由这种长文本的表格，需要根据输入调整行数
2. 不同次输入的文本不是一个runs。所以需要提示用户所有 应该连接在一起的文字 需要在同一次保存完成
4. 所有数字串需要在模板内使用完全不同的组合才能精准定位和替换
'''
